"""Document ingestion. Extracts raw text from PDF and DOCX files."""

import os
from pathlib import Path

def extract_text(file_path: str | Path) -> str:
    """Extract raw text from a BEP document (PDF or DOCX)."""
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = file_path.suffix.lower()
    
    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported document format: {ext}. Only .pdf and .docx are supported.")

def _extract_from_pdf(file_path: Path) -> str:
    import fitz  # pymupdf
    text = []
    with fitz.open(str(file_path)) as doc:
        for page in doc:
            text.append(page.get_text())
    return "\n".join(text)

def _extract_from_docx(file_path: Path) -> str:
    import docx
    doc = docx.Document(str(file_path))
    text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
            
    # Extract tables (often used in BEPs for Pset requirements)
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                text.append(" | ".join(row_data))
                
    return "\n".join(text)
